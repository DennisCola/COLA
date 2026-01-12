import streamlit as st
import pandas as pd
import google.generativeai as genai
from docx import Document
import json
import re

st.set_page_config(page_title="線控工作台", layout="wide")

# 設定 API
if "GEMINI_API_KEY" not in st.secrets:
    st.error("請在 Secrets 設定 API Key"); st.stop()

genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
model = genai.GenerativeModel('gemini-1.5-flash')

COLS = ["天數", "行程大點", "午餐", "晚餐", "有料門票", "旅館"]

st.title("🛡️ 線控行程「脫水」分類器")
st.write("---")

# 第一步：上傳與辨識
up = st.file_uploader("1. 請上傳 Word 行程表 (.docx)", type=["docx"])

if up:
    # 這裡加入一個緩存，避免重複扣 API 額度
    if 'raw_df' not in st.session_state or st.session_state.get('last_fn') != up.name:
        try:
            with st.spinner("AI 正在閱讀 Word 並過濾廢話..."):
                doc = Document(up)
                # 提取純文字並保持表格對應關係
                content = []
                for p in doc.paragraphs:
                    if p.text.strip(): content.append(p.text.strip())
                for tbl in doc.tables:
                    for row in tbl.rows:
                        row_data = [c.text.strip() for c in row.cells if c.text.strip()]
                        if row_data: content.append(" | ".join(dict.fromkeys(row_data)))
                
                full_text = "\n".join(content)
                
                # 脫水指令：強調只留精華
                prompt = f"""
                你是一位專業線控。請將這份 Word 行程『脫水』，濾掉所有推銷文字，只保留核心成本資訊。
                產出 JSON 列表，格式：{json.dumps(COLS, ensure_ascii=False)}。
                
                【脫水規則】：
                1. 『午/晚餐』：縮簡為餐食名稱（如：六菜一湯、米其林一星、自理）。
                2. 『有料門票』：僅列出需付費進入的景點，濾掉車窗參觀項目。
                3. 『旅館』：僅保留飯店名稱或星等。
                4. 必須嚴格按照天數排序。
                
                內容：
                {full_text[:6000]}
                """
                
                res = model.generate_content(prompt)
                match = re.search(r'\[.*\]', res.text, re.DOTALL)
                
                if match:
                    data = json.loads(match.group(0))
                    st.session_state.raw_df = pd.DataFrame(data).reindex(columns=COLS).fillna("").astype(str)
                    st.session_state.last_fn = up.name
                else:
                    st.error("AI 無法解析內容，請確認 Word 內容是否有誤。")
        except Exception as e:
            st.error(f"連線失敗：{e}")

# 第二步：展示表格（脫水結果）
if 'raw_df' in st.session_state:
    st.subheader("📍 第二步：確認脫水表格")
    st.info("請檢查 AI 抓取的內容是否正確，你可以直接點擊格子修改。這將作為報價的基礎。")
    
    # 使用 data_editor 讓使用者可以微調
    final_df = st.data_editor(
        st.session_state.raw_df, 
        use_container_width=True, 
        num_rows="dynamic",
        key="editor"
    )
    
    st.write("---")
    
    # 第三步：報價計算 (只有表格確認後才進行)
    st.subheader("💰 第三步：進入報價計算")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        ex_rate = st.number_input("歐元匯率", value=35.5)
    with col2:
        profit_margin = st.number_input("預期利潤 (%)", value=15)
    with col3:
        pax = st.number_input("成行人數", value=20)

    if st.button("🧮 計算總報價"):
        st.success(f"正在根據上述 {len(final_df)} 天行程計算成本...")
        # 這裡未來連動 Google Sheet 的單價資料庫
        st.info("此功能將連動 Google Sheet 成本資料庫（開發中）")
